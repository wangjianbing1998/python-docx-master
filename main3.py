# coding=utf-8
import argparse
import os
import uuid

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from get_data import get_df_huzhu
from utils import print_options

help_message = '三组-0,' + '九组-1,' + '十四组-2,' + '十一组-3,' + '四组-4,' + '十三组-5,' + '六组-6,' + '七组-7,' + '八组-8, 二组-9, 一组-10'
parser = argparse.ArgumentParser(description=help_message)
parser.add_argument('--file', '-f', default=7, type=int, help='the file to solve')
parser.add_argument('--testing', '-T', action='store_true', help='weather to testing')
args = parser.parse_args()
print_options(args, parser)

pd.set_option('expand_frame_repr', False)
pd.set_option('display.max_rows', 20)
pd.set_option('precision', 2)

testing = args.testing

sheet_names = [
    '三组',  # 0
    '九组',  # 1
    '十四组',  # 2
    '十一组',  # 3
    '四组',  # 4
    '十三组',  # 5
    '六组',  # 6
    '七组',  # 7
    '八组',  # 8
    '二组',  # 9
    '一组',  # 10
]

sheet_index = args.file
os.makedirs(f'{sheet_names[sheet_index]}/', exist_ok=True)

table_3 = Document(f"C:/Users/25536/Desktop/六组/表3.docx")
table_3.save(f'{sheet_names[sheet_index]}/附件3.docx')

excel_9 = pd.read_excel(r'../附件9.xlsx', sheet_name=sheet_index, dtype=object)
excel_9.rename(columns={'户主或与户主关系': '与户主关系'}, inplace=True)
excel_9.fillna('', inplace=True)

df, o_huzhu_dict = get_df_huzhu(sheet_names[sheet_index])

if testing:
    for i in range(len(excel_9) - len(df)):
        name = uuid.uuid1()
        df = df.append(df.iloc[-1, :], ignore_index=True, )
        df['成员姓名'] = name

for name in df['成员姓名'].values:
    if name not in excel_9['姓  名'].values:
        raise ValueError(f'name {name} not in excel 9')

for name in excel_9['姓  名'].values:
    if name not in df['成员姓名'].values:
        raise ValueError(f'name {name} not in df')

assert len(excel_9) == len(df), f'len(excel_9)={len(excel_9)}, len(df)={len(df)}'

huxuhao = {}
for huzhu_name, excel in o_huzhu_dict.items():
    huxuhao[huzhu_name] = excel.ix[0, '户序号']

excel_3 = excel_9[['姓  名', '身份证号码', '与户主关系', '性别']]
excel_3['户籍号'] = df['户籍号']
excel_3['户籍地址'] = df['户籍地址']
excel_3['存在状态'] = df['存在状态']
excel_3['该存在状态原因'] = df['该存在状态原因']
excel_3['兵役、独生子女、残疾人（等级）状况'] = df['兵役、独生子女、残疾人（等级）状况']
excel_3['土地（共有）使用权'] = df['土地（共有）使用权']
excel_3['保留型土地使用权'] = '√'
excel_3['承包经营权'] = df['承包经营权']
excel_3['集体资产管理权'] = '√'
excel_3['集体收益分配权'] = '√'
excel_3.rename(columns={'姓  名': '成员姓名'}, inplace=True)

col_name = excel_3.columns.tolist()
col_name.insert(0, '户主姓名')
col_name.insert(0, '序号')
excel_3 = excel_3.reindex(columns=col_name, fill_value='')

excel_3.loc[excel_3['与户主关系'].str.contains('户主'), '户主姓名'] = excel_3['成员姓名']

# excel_3.to_excel(f'{sheet_names[sheet_index]}/excel_3.xlsx', index=False)
print(f'load excel 3 successfully')

import copy

excel_3_df = copy.deepcopy(excel_3)
excel_3_df['联系电话'] = df['联系电话']
excel_3_df['兵役状况'] = df['兵役状况']
excel_3_df['现住地址'] = df['现住地址']
excel_3_df['婚姻状况'] = df['婚姻状况']
excel_3_df['文化程度'] = df['文化程度']

from docx.oxml.ns import qn

from utils import get_power, get_state, get_birth_day, get_huzhu_dict

excel_1 = excel_9[['姓  名', '性别', '与户主关系', '身份证号码']]
excel_1['户籍地址'] = df['户籍地址']
excel_1['现居住地址'] = df['现住地址']
excel_1['联系电话'] = df['联系电话']
excel_1['户主签字'] = ''

del excel_3_df['序号']
del excel_3_df['户主姓名']

huzhu_dict_3 = get_huzhu_dict(excel_3_df, huzhu_index=2, name_index=0)


def make_table_1(doc, excel_1):
    excel_data = (str(c) for c in excel_1.values.flatten().tolist())
    start = False
    end = False
    table1 = doc.tables[0]
    for row in table1.rows:
        if end:
            break
        for cell in row.cells:
            if end:
                break
            if cell.text == '户主情况':
                end = True
                break
            if start:
                try:
                    cell.text = next(excel_data)
                    # cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if cell.text == 'nan':
                        cell.text = ''
                except StopIteration:
                    cell.text = ''
            elif cell.text == '户主签字':
                start = True


def make_table_2(doc, home, ):
    # make table 2
    table2 = doc.tables[1]
    rows = table2.rows

    def set_value(i, j, s):
        cell = rows[i].cells[j]
        cell.text = s
        # rows[i].cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # cell=rows[i].cells[j]
        # for t in cell.paragraphs:
        #     t.clear()
        #
        # paragraphs = cell.paragraphs
        # for paragraph in paragraphs:
        #     for run in paragraph.runs:
        #         font = run.font
        #         font.size = Pt(30)
        #     paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #
        # my_paragraph = cell.paragraphs[0]
        # my_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run = my_paragraph.add_run(s)

    def set_align(i, j, type):
        cell = rows[i].cells[j]
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.paragraph_format.alignment = type

    def set_font_size(i, j, size=12):
        cell = rows[i].cells[j]
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(size)

    def tianbiao_my(person, n):
        base_row_index = n * 6
        r = 0 + base_row_index
        set_value(r, 2, person['成员姓名'])
        set_value(r, 5, person['性别'])
        set_value(r, 8, '汉')
        set_value(r, 10, get_birth_day(person['身份证号码']))
        set_font_size(r,10,9)
        set_value(r, 12, person['文化程度'])

        r = 1 + base_row_index
        set_value(r, 2, str(person['身份证号码']))
        set_value(r, 8, person['兵役状况'])  # 兵役情况

        # set_value(r, 10, phone_numbers[person['成员姓名']])
        set_value(r, 10, person['联系电话'])

        set_value(r, 12, person['与户主关系'])

        r = 2 + base_row_index
        set_value(r, 2, get_power(person[['土地（共有）使用权', '保留型土地使用权', '承包经营权', '集体资产管理权', '集体收益分配权']]))
        set_value(r, 7, get_state(person['存在状态']))

        set_align(r, 2, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_align(r, 7, WD_ALIGN_PARAGRAPH.LEFT)

        r = 3 + base_row_index
        # set_value(r, 7, get_state(person['存在状态']))
        set_value(r, 10, person['现住地址'])

        r = 4 + base_row_index
        # set_value(r, 2, get_power(person[['土地（共有）使用权', '保留型土地使用权', '承包经营权', '集体资产管理权', '集体收益分配权']]))
        set_value(r, 10, str(person['户籍号']))

        r = 5 + base_row_index
        set_value(r, 2, person['婚姻状况'])  # 婚姻状况

    huzhu = home.iloc[0, :]
    tianbiao_my(huzhu, 0)
    for p_index in range(1, len(home)):
        tianbiao_my(home.iloc[p_index, :], p_index)


def make_table_3(doc, excel):
    table1 = doc.tables[1]
    for _ in range(len(excel) - len(table1.rows)):
        table1.add_row()

    excel_data = (str(c) for c in excel.values.flatten().tolist())

    end = False
    for row in table1.rows:
        if end:
            break
        for cell in row.cells:
            if end:
                break
            try:
                cell.text = next(excel_data)
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except StopIteration:
                cell.text = ''
                end = True
                break


def update_paragraph(paragraph, s, bold=True, size=12):
    global run, font, font
    paragraph.text = s
    for run in paragraph.runs:
        font = run.font
        font.size = Pt(size)
        run.bold = bold


if __name__ == '__main__':

    huzhu_dict_1 = get_huzhu_dict(excel_1, huzhu_index=2, name_index=0)

    for huzhu_name, excel_1 in huzhu_dict_1.items():
        print(f'处理户主  {huzhu_name} ...')

        max_item = max(len(excel_1), 5)
        doc = Document(f"C:/Users/25536/Desktop/六组/附录1、2样表_{str(max_item)}.docx")

        update_paragraph(doc.paragraphs[2],
                         f'              马影 乡（镇、场）      石山  村（社区）   {"".join(sheet_names[sheet_index][:-1])}   组，户序号：{huxuhao[huzhu_name]}',
                         bold=True, size=12)

        update_paragraph(doc.paragraphs[10], f'马影 乡（镇、场）      石山 村（社区）   {"".join(sheet_names[sheet_index][:-1])} 组',
                         bold=True, size=12)
        doc.styles['Normal'].font.name = u'仿宋'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

        make_table_1(doc, excel_1)

        # show_excel(doc)
        # make table 2

        make_table_2(doc, huzhu_dict_3[huzhu_name], )
        os.makedirs(f'{sheet_names[sheet_index]}/附件1和附件2/', exist_ok=True)
        doc.save(f'{sheet_names[sheet_index]}/附件1和附件2/{huzhu_name}.docx')

        if testing:
            break

    # make table 3
    # excel_3 = pd.read_excel(f'{sheet_names[sheet_index]}/excel_3.xlsx', dtype=object)
    excel_3.fillna('', inplace=True)

    doc = Document(f"{sheet_names[sheet_index]}/附件3.docx")

    update_paragraph(doc.paragraphs[1], f'马影 乡（镇、场）      石山 村（社区）    {sheet_names[sheet_index]} {huxuhao[huzhu_name]}',
                     bold=True, size=12)

    make_table_3(doc, excel_3)
    doc.save(f"{sheet_names[sheet_index]}/附件3.docx")

    print(f'{sheet_names[sheet_index]} finished !')

    import shutil

    shutil.copy(r'C:\Users\25536\Desktop\六组\附件4.docx', f'{sheet_names[sheet_index]}/附件4.docx')
    shutil.copy(r'C:\Users\25536\Desktop\六组\附件5.docx', f'{sheet_names[sheet_index]}/附件5.docx')
    shutil.copy(r'C:\Users\25536\Desktop\六组\附件6.docx', f'{sheet_names[sheet_index]}/附件6.docx')
    shutil.copy(r'C:\Users\25536\Desktop\六组\附件7.docx', f'{sheet_names[sheet_index]}/附件7.docx')
    shutil.copy(r'C:\Users\25536\Desktop\六组\附件8.docx', f'{sheet_names[sheet_index]}/附件8.docx')
    shutil.copy(r'C:\Users\25536\Desktop\六组\附件10.docx', f'{sheet_names[sheet_index]}/附件10.docx')
    shutil.copy(
        f'C:/Users/25536/Documents/WeChat Files/wjb15071291411/FileStorage/File/2020-03/经济组织成员各组公示表/{sheet_names[sheet_index]}.docx',
        f'{sheet_names[sheet_index]}/附件9.docx')
